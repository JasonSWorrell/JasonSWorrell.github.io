from annotated_types import doc
import gradio as gr
import torch
from pydub import AudioSegment
from docx import Document
from App_Model import Summarizer, NoteTaker, Transcriber

class AppGradioUI:
    def __init__(self):

        with gr.Blocks(theme=gr.themes.Soft()) as self.demo:
            gr.Markdown("# 📝 AI Note-Taker Assistant")
            gr.Markdown("Record a meeting or upload an audio file to generate instant transcripts and action items.")
            
            with gr.Row():
                with gr.Column():
                    audio = gr.Audio(sources=["microphone", "upload"], type="filepath", label="Input Audio Source")
                    submit_btn = gr.Button("Generate Notes", variant="primary")
                    filename = gr.Textbox(label="File Name", placeholder="Name of file", lines=1, value="AI_Notes")
                    PDF_btn = gr.Button("Generate Notes and PDF", variant="primary")

                    
                with gr.Column():
                    transcript_output = gr.Textbox(label="Full Transcript", lines=8)
                    summary_output = gr.Textbox(label="AI Summary & Action Items", lines=6)
                    notes_output = gr.Textbox(label="AI Notes", lines=12)

                    
            submit_btn.click(
                fn=self.process_audio,
                inputs=[audio],
                outputs=[transcript_output, summary_output, notes_output]
            )

            PDF_btn.click(
                fn=self.doc,
                inputs=[filename, audio],
                outputs=[transcript_output, summary_output, notes_output]
            )

    def process_audio(self,audio_path): # audio_path is the input audio, may need to be a few seconds or few minutes at a time. 
        transcriber = Transcriber("openai/whisper-small", "cuda" if torch.cuda.is_available() else "cpu")
        summarizer = Summarizer("knkarthick/MEETING_SUMMARY", "cuda" if torch.cuda.is_available() else "cpu")   
        note_taker = NoteTaker("meta-llama/Llama-3.1-8B-Instruct", "cuda" if torch.cuda.is_available() else "cpu")

        transcript = transcriber.process_audio(audio_path)
        summary = summarizer.process_transcription(transcript)
        notes = note_taker.process_transcription(transcript)
        return transcript, summary, notes

    def doc(self, filename, audio_path):
        transcription, summary, txt = self.process_audio(audio_path)
        document = Document()
        document.add_heading(f'{filename}', 0)
        document.add_paragraph(f"Summary: {summary}")
        document.add_paragraph(txt)
        document.save(f'Notes/{filename}.docx')
        return transcription, summary, txt

class AppKinterUI:
    def __init__(self, SUMMARIZER, NOTE_TAKER, TRANSCRIBER):
        self.summarizer = SUMMARIZER
        self.note_taker = NOTE_TAKER
        self.transcriber = TRANSCRIBER

from UI import GradioUI, KinterUI
import torch
from pydub import AudioSegment
from docx import Document
from App_Model import Summarizer, NoteTaker, Transcriber
import sys


def process_audio(audio_path): # audio_path is the input audio, may need to be a few seconds or few minutes at a time. 
    transcript = Transcriber.process_transcription(audio_path)
    summary = Summarizer.process_transcription(transcript)
    notes = NoteTaker.process_transcription(transcript)
    return transcript, summary, notes

def doc(filename,audio_path):
    transcription, summary, txt = process_audio(audio_path)
    document = Document()
    document.add_heading(f'{filename}', 0)
    document.add_paragraph(f"Summary: {summary}")
    document.add_paragraph(txt)
    document.save(f'Notes/{filename}.docx')
    return transcription, summary, txt

def main():
    if len(sys.argv) > 1:
        return "Acceptable commandline arguments are kinter or gradio. Otherwise please run the script without any command line arguments and the Gradio UI will launch automatically."
    else:
        if sys.argv == 1:            
                
            transcriber = Transcriber("openai/whisper-small", 
                                        "cuda" if torch.cuda.is_available() else "cpu")
            summarizer = Summarizer("knkarthick/MEETING_SUMMARY", 
                                    "cuda" if torch.cuda.is_available() else "cpu")   
            note_taker = NoteTaker("meta-llama/Llama-3.1-8B-Instruct", 
                                    "cuda" if torch.cuda.is_available() else "cpu")  
            if sys.argv[0] == "gradio":
                app = GradioUI(summarizer, note_taker, transcriber)
            if sys.argv[0] == "kinter":
                app = KinterUI(summarizer, note_taker, transcriber)
            else:
                print("The provided GUI argument was invalid. ")
                app = GradioUI(summarizer, note_taker, transcriber)
            app.launch()




if __name__ == "__main__":
    main()
